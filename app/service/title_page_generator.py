import os
import tempfile
import re

from uuid import uuid4

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


MINISTRY_NAME = "МИНОБРНАУКИ РОССИИ"
UNIVERSITY_FULL_NAME = "Федеральное государственное бюджетное образовательное учреждение высшего образования"
UNIVERSITY_SHORT_NAME = "«Юго-Западный государственный университет» (ЮЗГУ)"


def set_font(run, size=16, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def add_paragraph(
    container,
    text: str = "",
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    size: int = 16,
    bold: bool = False,
    space_before: int = 0,
    space_after: int = 0,
    line_spacing: float = 1.0,
):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_empty_paragraphs(container, count: int, size: int = 16):
    for _ in range(count):
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run("")
        set_font(run, size=size)


def normalize_department_name(name: str) -> str:
    name = (name or "").strip()
    lower_name = name.lower()
    if lower_name.startswith("кафедра "):
        return name[8:].strip()
    return name


def set_footer_text(footer, text: str, size: int = 16):
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)
    set_font(run, size=size, bold=False)


def set_first_page_footer(section, text: str, size: int = 16):
    section.different_first_page_header_footer = True
    set_footer_text(section.first_page_footer, text, size=size)
    set_footer_text(section.footer, text, size=size)


def clear_section_footer(section):
    section.different_first_page_header_footer = True
    clear_paragraph(section.first_page_footer.paragraphs[0])
    clear_paragraph(section.footer.paragraphs[0])


def format_tutorial_title(title: str) -> str:
    title = (title or "").strip()
    if not title:
        return title
    return title[:1].upper() + title[1:].lower()


def ensure_a_value_prefix(a_value: str) -> str:
    value = (a_value or "").strip()
    if not value:
        return "А __"

    upper_value = value.upper()
    if upper_value.startswith("А "):
        return value
    if upper_value.startswith("A "):
        return f"А {value[2:].strip()}"
    if upper_value == "А" or upper_value == "A":
        return "А __"

    return f"А {value}"


def format_author_name_for_biblio(author_name: str) -> str:
    """
    Поддерживает:
    - Аникина Елена Ивановна -> Е. И. Аникина
    - Аникина Е. И. -> Е. И. Аникина
    - Аникина Е.И. -> Е. И. Аникина
    - Е. И. Аникина -> Е. И. Аникина
    - Е.И. Аникина -> Е. И. Аникина
    """
    text = (author_name or "").strip()
    if not text:
        return text

    text = re.sub(r"\s+", " ", text)

    # случай: Е. И. Аникина
    m = re.match(r"^([А-ЯA-Z])\.\s*([А-ЯA-Z])\.\s*([А-Яа-яA-Za-z\-]+)$", text)
    if m:
        n1, n2, surname = m.groups()
        return f"{n1}. {n2}. {surname}"

    # случай: Е.И. Аникина
    m = re.match(r"^([А-ЯA-Z])\.([А-ЯA-Z])\.\s*([А-Яа-яA-Za-z\-]+)$", text)
    if m:
        n1, n2, surname = m.groups()
        return f"{n1}. {n2}. {surname}"

    parts = text.split()

    # случай: Аникина Елена Ивановна
    if len(parts) == 3 and all(len(p) > 1 for p in parts):
        surname, name, patronymic = parts
        return f"{name[0]}. {patronymic[0]}. {surname}"

    # случай: Аникина Е. И.
    if len(parts) == 3 and len(parts[0]) > 1:
        surname = parts[0]
        p2 = parts[1].replace(".", "")
        p3 = parts[2].replace(".", "")
        if len(p2) == 1 and len(p3) == 1:
            return f"{p2}. {p3}. {surname}"

    return text

def generate_title_page_docx(
    manual_title: str,
    discipline_name: str,
    audience: str,
    direction_code: str,
    direction_name: str,
    department_name: str,
    city: str,
    year: int,
    udk: str,
    compiler_name: str,
    reviewer_name: str,
    reviewer_degree: str,
    description: str,
) -> str:
    output_dir = os.path.join(tempfile.gettempdir(), "title_pages")
    os.makedirs(output_dir, exist_ok=True)

    filename = f"title_page_{uuid4().hex}.docx"
    output_path = os.path.join(output_dir, filename)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.25)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.3)
    section.footer_distance = Cm(1.3)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(16)

    clean_department_name = normalize_department_name(department_name)



    add_paragraph(doc, MINISTRY_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, UNIVERSITY_FULL_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, UNIVERSITY_SHORT_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(
        doc,
        f"Кафедра {clean_department_name}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=16,
    )

    add_empty_paragraphs(doc, 1)

    add_paragraph(
        doc,
        "УТВЕРЖДАЮ",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=16,
        bold=True,
    )
    add_paragraph(
        doc,
        "Заведующий кафедрой __________",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=16,
    )
    add_paragraph(
        doc,
        "______________________________",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=16,
    )
    add_paragraph(
        doc,
        '"____" __________ 20____ г.',
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=16,
    )

    add_empty_paragraphs(doc, 4)

    add_paragraph(
        doc,
        manual_title.upper(),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=16,
        bold=True,
        space_after=6,
    )

    add_paragraph(
        doc,
        (
            f"Методические указания для выполнения лабораторной работы "
            f"по дисциплине «{discipline_name}» для {audience} "
            f"направления подготовки {direction_code} {direction_name}"
        ),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=16,
        line_spacing=1.0,
    )

    add_paragraph(
        doc,
        f"{city} - {year}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=16,
        space_before=210,
    )

    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.top_margin = Cm(3)
    new_section.bottom_margin = Cm(2.25)
    new_section.left_margin = Cm(2.2)
    new_section.right_margin = Cm(2.3)
    new_section.footer_distance = Cm(1.3)
    clear_section_footer(new_section)

    add_paragraph(doc, f"УДК  {udk}", size=16, space_after=0)
    add_paragraph(doc, f"Составитель:  {compiler_name}", size=16, space_after=6)

    add_paragraph(doc, "Рецензент", size=16, space_after=0)
    add_paragraph(
        doc,
        f"{reviewer_degree}  {reviewer_name}",
        size=16,
        space_after=6,
    )

    add_paragraph(
        doc,
        (
            f"{manual_title}: методические указания для выполнения  лабораторной работы "
            f"по дисциплине «{discipline_name}» для {audience} "
            f"направления подготовки {direction_code} {direction_name}/ "
            f"Юго-Зап. гос. ун-т; сост. {compiler_name}. "
            f"{city}, {year}. 19 с."
        ),
        size=16,
        line_spacing=1.0,
        space_after=6,
    )

    add_paragraph(
        doc,
        (
            f"Составлены в соответствии с федеральным государственным образовательным стандартом "
            f"высшего образования направления подготовки {direction_code} {direction_name} "
            f"и на основании учебного плана направления подготовки {direction_code} {direction_name}."
        ),
        size=16,
        line_spacing=1.0,
        space_after=6,
    )

    add_paragraph(
        doc,
        description,
        size=16,
        line_spacing=1.0,
        space_after=6,
    )

    add_paragraph(
        doc,
        (
            f"Предназначены для {audience}, обучающихся направления подготовки "
            f"{direction_code} {direction_name} "
            f"(профиль «Разработка программноинформационных систем») всех форм обучения."
        ),
        size=16,
        line_spacing=1.0,
        space_after=8,
    )

    add_paragraph(
        doc,
        "Текст печатается в авторской редакции",
        size=16,
        space_after=12,
    )

    add_paragraph(
        doc,
        (
            "Подписано в печать __________. Формат 60x84 1/16. "
            "Усл. печ. л. ____. Уч.-изд. л. 2,0. "
            "Тираж ___ экз. Заказ ________. Бесплатно."
        ),
        size=16,
        line_spacing=1.0,
        space_after=6,
    )

    add_paragraph(
        doc,
        "Юго-Западный государственный университет.",
        size=16,
        line_spacing=1.0,
        space_after=0,
    )

    add_paragraph(
        doc,
        "305040, г. Курск, ул. 50 лет Октября, 94.",
        size=16,
        line_spacing=1.0,
        space_after=0,
    )

    doc.save(output_path)
    return output_path


def generate_tutorial_title_page_docx(
    author_name: str,
    tutorial_title: str,
    city: str,
    year: int,
    reviewers: list,
    a_value: str,
    isbn: str,
    directions: list,
    udk: str,
    bbk: str,
    description: str,
) -> str:
    output_dir = os.path.join(tempfile.gettempdir(), "title_pages")
    os.makedirs(output_dir, exist_ok=True)

    filename = f"tutorial_title_page_{uuid4().hex}.docx"
    output_path = os.path.join(output_dir, filename)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.25)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.3)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(16)

    title_for_biblio = format_tutorial_title(tutorial_title)
    a_value_formatted = ensure_a_value_prefix(a_value)
    author_name_biblio = format_author_name_for_biblio(author_name)

    # 1 страница
    add_paragraph(doc, "МИНОБРНАУКИ РОССИИ", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, "Федеральное государственное бюджетное", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, "образовательное учреждение высшего", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, "образования", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, "«Юго-Западный государственный университет»", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, "(ЮЗГУ)", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)

    add_empty_paragraphs(doc, 6)

    add_paragraph(doc, author_name_biblio, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_empty_paragraphs(doc, 2)

    add_paragraph(doc, tutorial_title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    add_empty_paragraphs(doc, 1)
    add_paragraph(doc, "Учебное пособие", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_empty_paragraphs(doc, 6)

    add_paragraph(
        doc,
        "Утверждено Учебно-методическим составом",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=16,
    )
    add_paragraph(
        doc,
        "Юго-Западного государственного университета",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=16,
    )

    # Курск 2026 внизу первой страницы не через footer, а через отступ
    add_paragraph(
        doc,
        f"{city} {year}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=16,
        space_before=140,
    )

    # 2 страница
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.top_margin = Cm(3)
    new_section.bottom_margin = Cm(2.25)
    new_section.left_margin = Cm(2.2)
    new_section.right_margin = Cm(2.3)
    clear_section_footer(new_section)

    reviewers_title = "Рецензенты:" if len(reviewers) > 1 else "Рецензент:"
    add_paragraph(doc, reviewers_title, size=16, space_after=4)

    for reviewer in reviewers:
        text = f"{reviewer.degree_position}  {reviewer.fio}"
        if text.strip().lower().startswith("рецензенты:"):
            text = text.split(":", 1)[1].strip()
        if text.strip().lower().startswith("рецензент:"):
            text = text.split(":", 1)[1].strip()

        add_paragraph(
            doc,
            text,
            size=16,
            space_after=4,
        )

    add_empty_paragraphs(doc, 1)

    add_paragraph(doc, author_name_biblio, size=16)
    add_empty_paragraphs(doc, 1)

    add_paragraph(
        doc,
        f"{a_value_formatted}    {title_for_biblio}: учеб. пособие / {author_name_biblio};",
        size=16,
    )
    add_paragraph(
        doc,
        f"Юго-Зап. гос. ун-т. – {city}, {year}. – ____ с. – Библиогр.: с. ___.",
        size=16,
    )
    add_paragraph(doc, f"ISBN {isbn}", size=16, space_after=8)

    add_paragraph(doc, description, size=16, line_spacing=1.0, space_after=8)

    if len(directions) == 1:
        directions_text = f"{directions[0].code} «{directions[0].faculty_name}»"
    else:
        directions_text = ", ".join(
            [f"{item.code} «{item.faculty_name}»" for item in directions[:-1]]
        ) + f" и {directions[-1].code} «{directions[-1].faculty_name}»"

    add_paragraph(
        doc,
        f"Предназначено для студентов и магистрантов укрупненных групп направлений подготовки {directions_text}.",
        size=16,
        line_spacing=1.0,
        space_after=8,
    )

    add_paragraph(doc, f"УДК {udk}", size=16)
    add_paragraph(doc, f"ББК {bbk}", size=16, space_after=8)

    add_paragraph(doc, f"ISBN {isbn}", size=16, space_after=8)

    add_paragraph(doc, f"© Юго-Западный государственный университет, {year}", size=16)
    add_paragraph(doc, f"© {author_name_biblio}, {year}", size=16)

    doc.save(output_path)
    return output_path


def generate_monograph_title_page_docx(
        authors: list,
        monograph_title: str,
        city: str,
        year: int,
        udk: str,
        bbk: str,
        isbn: str,
        description: str,
) -> str:
    """
    Генерация титульного листа монографии
    """
    output_dir = os.path.join(tempfile.gettempdir(), "title_pages")
    os.makedirs(output_dir, exist_ok=True)

    filename = f"monograph_title_page_{uuid4().hex}.docx"
    output_path = os.path.join(output_dir, filename)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.25)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.3)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(16)

    authors_fio = [author.fio if hasattr(author, 'fio') else author for author in authors]

    if len(authors_fio) == 1:
        authors_biblio = authors_fio[0]
    else:
        authors_biblio = f"{', '.join(authors_fio[:-1])} и {authors_fio[-1]}"

    # ========== 1 СТРАНИЦА ==========
    add_paragraph(doc, "МИНИСТЕРСТВО ОБРАЗОВАНИЯ РФ", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, "Государственное образовательное учреждение высшего профессионального образования",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, "«Юго-Западный государственный университет»", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)

    add_empty_paragraphs(doc, 4)

    for author in authors_fio:
        add_paragraph(doc, author, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)

    add_empty_paragraphs(doc, 2)

    add_paragraph(doc, monograph_title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)

    add_empty_paragraphs(doc, 1)

    add_paragraph(doc, "Монография", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)

    add_empty_paragraphs(doc, 6)

    add_paragraph(doc, f"{city} -- {year}", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)

    # ========== 2 СТРАНИЦА ==========
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.top_margin = Cm(3)
    new_section.bottom_margin = Cm(2.25)
    new_section.left_margin = Cm(2.2)
    new_section.right_margin = Cm(2.3)
    clear_section_footer(new_section)

    add_paragraph(doc, f"УДК {udk}", size=16, space_after=0)
    add_paragraph(doc, f"ББК {bbk}", size=16, space_after=12)

    add_paragraph(doc, "Рецензенты:", size=16, space_after=6)

    add_empty_paragraphs(doc, 1)

    biblio_text = f"{authors_biblio}. {monograph_title}: монография / {authors_biblio}; Юго-Западный гос. ун-т. – {city}, {year}. – ____ с. : ил. ____, табл. ____, Библиогр.: с. _____"
    add_paragraph(doc, biblio_text, size=16, line_spacing=1.0, space_after=6)

    add_paragraph(doc, f"ISBN {isbn}", size=16, space_after=12)

    add_paragraph(doc, description, size=16, line_spacing=1.0, space_after=12)

    add_paragraph(doc, f"УДК {udk}", size=16, space_after=0)
    add_paragraph(doc, f"ББК {bbk}", size=16, space_after=12)

    add_paragraph(doc, f"ISBN {isbn}", size=16, space_after=6)
    add_paragraph(doc, f"© Юго-Западный государственный университет, {year}", size=16, space_after=0)
    add_paragraph(doc, f"© {authors_biblio}, {year}", size=16)

    doc.save(output_path)
    return output_path